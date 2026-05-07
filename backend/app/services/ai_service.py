import os
import json
import base64
import requests
from typing import List, Dict, Any
import fitz  # PyMuPDF
from ..config import settings
import PIL.Image

# Conditionally import new google.genai SDK (replaces deprecated google.generativeai)
try:
    from google import genai
    from google.genai import types as genai_types
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False

# Conditionally import python-docx
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def extract_text_from_pdf(pdf_path: str) -> str:
    """Extracts text from a PDF file using PyMuPDF."""
    text = ""
    try:
        with fitz.open(pdf_path) as doc:
            for page in doc:
                text += page.get_text()
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
    return text

def extract_text_from_docx(docx_path: str) -> str:
    """Extracts text from a DOCX file using python-docx."""
    if not DOCX_AVAILABLE:
        print("python-docx not installed. Skipping DOCX extraction.")
        return ""
    text = ""
    try:
        doc = DocxDocument(docx_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
            text += "\n"
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
    return text

def parse_json_response(text: str) -> Dict[str, Any]:
    json_str = text.strip()
    if "```json" in json_str:
        json_str = json_str.split("```json")[1].split("```")[0].strip()
    elif "```" in json_str:
        json_str = json_str.split("```")[1].split("```")[0].strip()
    try:
        return json.loads(json_str)
    except json.JSONDecodeError as e:
        print(f"JSON Parsing Error: {e}\nRaw Output:\n{text}")
        return {}

def call_gemini(prompt_parts: List[Any]) -> Dict[str, Any]:
    """Calls Gemini model and expects a JSON response (google.genai SDK)."""
    if not GEMINI_AVAILABLE or not settings.GEMINI_API_KEY:
        raise ValueError("Gemini is not configured or missing API key.")

    client = genai.Client(api_key=settings.GEMINI_API_KEY)
    print("Routing request to Primary Engine (Google Gemini)...")

    # Separate text strings from PIL images for the new SDK
    contents = []
    for part in prompt_parts:
        if isinstance(part, str):
            contents.append(part)
        else:
            # PIL image — convert to bytes for the new SDK
            import io
            buf = io.BytesIO()
            part.save(buf, format="PNG")
            contents.append(
                genai_types.Part.from_bytes(data=buf.getvalue(), mime_type="image/png")
            )

    response = client.models.generate_content(
        model=settings.GEMINI_MODEL,
        contents=contents
    )
    return parse_json_response(response.text)

def call_ollama(prompt: str, model: str, images: List[str] = None) -> Dict[str, Any]:
    """Calls local Ollama model and expects a JSON response."""
    url = f"{settings.OLLAMA_HOST}/api/generate"
    payload = {
        "model": model,
        "prompt": prompt,
        "stream": False,
        "format": "json"
    }
    if images:
        payload["images"] = images
    
    print(f"Routing request to Fallback Engine (Local Ollama: {model})...")
    response = requests.post(url, json=payload, timeout=300) # Local LLMs can be slow
    response.raise_for_status()
    result = response.json()
    return json.loads(result.get("response", "{}"))

def extract_criteria_from_tender(file_path: str, preferred_engine: str = "auto") -> Dict[str, Any]:
    """
    Parses a tender document and extracts eligibility criteria.
    Tries preferred engine, falls back to others on failure.
    """
    text = extract_text_from_pdf(file_path)
    if not text:
        return {}

    prompt = f"""
    You are an expert CRPF Procurement Auditor. Analyze the tender and extract:
    1. ELIGIBILITY CRITERIA: Technical, Financial, Compliance. 
       - For Financial: Look for "Average Annual Turnover" requirements.
    2. DOCUMENT CHECKLIST: List every mandatory certification or document requested (e.g., GST, ISO 9001, PAN, MSME, Experience Certificate).

    Return ONLY JSON with this format:
    {{
      "criteria": [
        {{"category": "Financial", "name": "Avg Turnover", "threshold": "> ₹5 Cr (3yr Avg)", "mandatory": true, "source_page": 1, "source_text": "..."}}
      ],
      "required_docs": ["GST Registration", "ISO 9001:2015", "Audited Balance Sheets (3 yrs)"]
    }}

    TENDER TEXT:
    {text}
    """
    
    try:
        if preferred_engine == "ollama":
            res = call_ollama(prompt, settings.OLLAMA_TEXT_MODEL)
            res["_engine"] = "ollama"
            return res
        
        if preferred_engine == "gemini":
            res = call_gemini([prompt])
            res["_engine"] = "gemini"
            return res

        # Auto logic
        try:
            res = call_gemini([prompt])
            res["_engine"] = "gemini"
            return res
        except Exception as gemini_err:
            print(f"Gemini failed in auto-mode: {gemini_err}. Falling back to Ollama...")
            res = call_ollama(prompt, settings.OLLAMA_TEXT_MODEL)
            res["_engine"] = "ollama"
            return res

    except Exception as e:
        print(f"AI Engine Failed (preferred: {preferred_engine}): {e}")
        print("Falling back to Simulated Extraction (No AI Engines active)")
        return {
            "_engine": "simulated",
            "criteria": [
                {"category": "Financial", "name": "Annual Turnover", "threshold": "> ₹5 Cr", "mandatory": True, "source_page": 2},
                {"category": "Technical", "name": "Experience in Civil Works", "threshold": "3+ Projects", "mandatory": True, "source_page": 5},
                {"category": "Compliance", "name": "ISO 9001:2015", "threshold": "Valid Certificate", "mandatory": False, "source_page": 8}
            ],
            "required_docs": ["Financial Audit FY23", "Project Completion Certificates", "ISO Certificate"]
        }

def evaluate_bidder_against_criteria(bidder_files: List[str], criteria: List[Dict[str, Any]], preferred_engine: str = "auto") -> Dict[str, Any]:
    """
    Evaluates a bidder's documents against tender criteria.
    Tries preferred engine (Gemini/Ollama), falls back on failure.
    """
    criteria_json = json.dumps(criteria)

    base_prompt = (
        "You are a STRICT government procurement auditor. Your ONLY job is to verify whether "
        "the submitted documents EXPLICITLY contain evidence for each criterion.\n\n"
        "STRICT RULES — YOU MUST FOLLOW THESE WITHOUT EXCEPTION:\n"
        "1. NEVER assume, infer, or guess that a criterion is met. You MUST find EXPLICIT text/data in the documents.\n"
        "2. If NO direct evidence is found for a criterion, the verdict MUST be 'ineligible' — NOT 'review' and NOT 'eligible'.\n"
        "3. If a document is clearly irrelevant (e.g. a resume, a news article, an unrelated contract), ALL criteria must be 'ineligible'.\n"
        "4. 'review' is ONLY for cases where partial/ambiguous evidence exists (e.g. blurry image, conflicting numbers).\n"
        "5. NEVER mark a criterion 'eligible' based on absence of contradicting information.\n"
        "6. Extracted values MUST come directly from the document text. Do NOT fabricate numbers.\n\n"
        f"CRITERIA TO CHECK:\n{criteria_json}\n\nBIDDER DOCUMENTS:"
    )

    gemini_parts = [base_prompt]
    ollama_prompt = base_prompt
    ollama_images_base64 = []
    has_images = False

    all_text_content = ""  # Accumulate all extracted text for relevance check

    for file_path in bidder_files:
        ext = file_path.lower().split('.')[-1]
        if ext == 'pdf':
            text = extract_text_from_pdf(file_path)
            all_text_content += text
            gemini_parts.append(f"\n--- PDF DOCUMENT: {os.path.basename(file_path)} ---\n{text}")
            ollama_prompt += f"\n--- PDF DOCUMENT: {os.path.basename(file_path)} ---\n{text}"
        elif ext in ['docx', 'doc']:
            text = extract_text_from_docx(file_path)
            all_text_content += text
            gemini_parts.append(f"\n--- WORD DOCUMENT: {os.path.basename(file_path)} ---\n{text}")
            ollama_prompt += f"\n--- WORD DOCUMENT: {os.path.basename(file_path)} ---\n{text}"
        elif ext in ['jpg', 'jpeg', 'png']:
            try:
                has_images = True
                with PIL.Image.open(file_path) as img:
                    gemini_parts.append(f"\n--- IMAGE DOCUMENT: {os.path.basename(file_path)} ---")
                    gemini_parts.append(img)

                ollama_prompt += f"\n--- IMAGE DOCUMENT: {os.path.basename(file_path)} (See attached base64 image) ---"
                with open(file_path, "rb") as image_file:
                    encoded_string = base64.b64encode(image_file.read()).decode('utf-8')
                    ollama_images_base64.append(encoded_string)
            except Exception as e:
                print(f"Error loading image {file_path}: {e}")
        else:
            skipped_msg = f"\n--- FILE: {os.path.basename(file_path)} (Skipped extraction) ---"
            gemini_parts.append(skipped_msg)
            ollama_prompt += skipped_msg

    # --- PRE-FLIGHT RELEVANCE CHECK ---
    # If extracted text is too short or contains no procurement keywords, the document
    # is clearly irrelevant. Immediately return all-ineligible without calling the AI.
    PROCUREMENT_KEYWORDS = [
        "turnover", "revenue", "balance sheet", "profit", "loss", "gst", "pan",
        "iso", "certificate", "crore", "lakh", "registration", "experience",
        "project", "contract", "tender", "bid", "supplier", "manufacturer",
        "financial", "audit", "annual report", "compliance", "license", "msme"
    ]
    text_lower = all_text_content.lower()
    keyword_hits = sum(1 for kw in PROCUREMENT_KEYWORDS if kw in text_lower)
    MIN_TEXT_LENGTH = 100      # Fewer than 100 chars → almost certainly empty/corrupt
    MIN_KEYWORD_HITS = 2       # Fewer than 2 procurement keywords → clearly irrelevant

    if not has_images and (len(all_text_content.strip()) < MIN_TEXT_LENGTH or keyword_hits < MIN_KEYWORD_HITS):
        print(
            f"PRE-FLIGHT REJECTION: Document text too short ({len(all_text_content.strip())} chars) "
            f"or missing procurement keywords ({keyword_hits} hits). Marking all criteria ineligible."
        )
        evaluations = []
        checklist = {}
        for crit in criteria:
            checklist[crit.get("name", "Unknown")] = "missing"
        for crit in criteria:
            evaluations.append({
                "criterion_id": crit["id"],
                "verdict": "ineligible",
                "confidence": "high",
                "match_type": "none",
                "extracted_value": "Not found in submitted documents",
                "reasoning": (
                    f"REJECTED: The uploaded file(s) do not appear to be valid procurement documents. "
                    f"No relevant financial, technical, or compliance evidence was detected. "
                    f"Text extracted: {len(all_text_content.strip())} characters, "
                    f"procurement keyword hits: {keyword_hits}. "
                    f"Please upload the correct documents for criterion '{crit.get('name', 'Unknown')}'."
                ),
                "source_page": 0,
                "source_document": os.path.basename(bidder_files[0]) if bidder_files else "Unknown",
                "evidence_snippet": "N/A — irrelevant document",
                "action_required": (
                    f"Upload a valid document proving '{crit.get('name', 'Unknown')}' "
                    f"(e.g. audited balance sheet, GST certificate, ISO certificate)."
                )
            })
        return {"evaluations": evaluations, "checklist": checklist, "_pre_flight_rejected": True}

    prompt_suffix = """
    For each criterion, decide if the bidder is: 'eligible', 'ineligible', or 'review'.

    STRICT VERDICT RULES (MUST be followed — no exceptions):
    - 'eligible'   → ONLY if EXPLICIT, DIRECT evidence is found in the document. Quote the exact text.
    - 'ineligible' → If NO evidence exists OR the document is unrelated to the criterion.
    - 'review'     → ONLY if partial/ambiguous evidence exists (blurry image, conflicting numbers).

    SPECIAL LOGIC:
    - FINANCIAL AVERAGING: If the threshold is an average (e.g. 3 years), extract the value for each of the 3 years and calculate the average yourself. Show this calculation in the reasoning.
    - ZERO TOLERANCE: Do NOT mark 'eligible' based on assumption, inference, or because nothing contradicts it.
    - AMBIGUITY: If you choose 'review', explicitly state IF it's due to 'Image Quality', 'Missing Evidence', or 'Conflicting Information'.

    DOCUMENT CHECKLIST:
    - Based on the BIDDER DOCUMENTS, check if each item in the 'required_docs' list is EXPLICITLY present.
    - Mark 'missing' if the document type is not found or not clearly identifiable.

    Return ONLY a JSON object:
    {
      "evaluations": [
         {
           "criterion_id": "...",
           "verdict": "eligible" | "ineligible" | "review",
           "reasoning": "[Must quote exact evidence OR state exactly why it is ineligible]",
           "extracted_value": "[Exact value from document, or 'Not found']",
           "confidence": "high/medium/low",
           "source_page": 1,
           "source_document": "[Filename]",
           "evidence_snippet": "[The exact text snippet found, or 'No evidence found']",
           "action_required": "[Specific next step for human if verdict is 'review' or 'ineligible']"
         }
      ],
      "checklist": {
         "Document Name": "found" | "missing"
      }
    }
    """
    gemini_parts.append(prompt_suffix)
    ollama_prompt += prompt_suffix

    try:
        model_to_use = settings.OLLAMA_VISION_MODEL if has_images else settings.OLLAMA_TEXT_MODEL
        
        if preferred_engine == "ollama":
            res = call_ollama(ollama_prompt, model_to_use, ollama_images_base64 if has_images else None)
            res["_engine"] = "ollama"
            return res
        
        if preferred_engine == "gemini":
            res = call_gemini(gemini_parts)
            res["_engine"] = "gemini"
            return res

        # Auto logic
        try:
            res = call_gemini(gemini_parts)
            res["_engine"] = "gemini"
            return res
        except Exception as gemini_err:
            print(f"Gemini failed in auto-mode: {gemini_err}. Falling back to Ollama...")
            res = call_ollama(ollama_prompt, model_to_use, ollama_images_base64 if has_images else None)
            res["_engine"] = "ollama"
            return res

    except Exception as e:
        print(f"AI Engine Failed (preferred: {preferred_engine}): {e}")
        # Both AI engines are offline — do NOT fake a score.
        # Return 'review' verdicts so a human officer must verify manually.
        print("Both AI engines offline. Returning 'needs human review' result — no score will be assigned.")

        evaluations = []
        checklist = {}
        for crit in criteria:
            checklist[crit['name']] = "unverified"
            evaluations.append({
                "criterion_id": crit['id'],
                "verdict": "review",
                "confidence": "low",
                "match_type": "none",
                "extracted_value": "AI engines offline — manual review required",
                "reasoning": f"Automated evaluation unavailable for '{crit['name']}'. No AI engine (Gemini or Ollama) is currently active. A procurement officer must review this criterion manually.",
                "source_page": 1,
                "source_document": os.path.basename(bidder_files[0]) if bidder_files else "Unknown",
                "evidence_snippet": "N/A — AI offline",
                "action_required": "Manual review required: AI engine unavailable at time of upload."
            })
        # _simulated flag tells the backend not to compute a fake match score
        return {"evaluations": evaluations, "checklist": checklist, "_simulated": True, "_engine": "simulated"}
